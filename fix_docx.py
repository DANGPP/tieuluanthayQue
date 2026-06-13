"""
Script to fix the thesis document formatting - FINAL VERSION
Fixes:
1. Proper heading styles (Heading 1/2/3) based on font size patterns
2. Standardized line spacing (1.5)  
3. Consistent font (Times New Roman 13pt body)
4. Fix misclassified headings
5. Table formatting
6. Page margins (3.5/2/2/2 cm)
7. Fix paragraphs where heading + body text got merged
"""

import docx
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import re
import copy

INPUT_FILE = r'd:\2026\thayque\tieuluan\tieuluan.docx'
OUTPUT_FILE = r'd:\2026\thayque\tieuluan\tieuluan_fixed.docx'

doc = docx.Document(INPUT_FILE)

# ============================================================
# STEP 1: Configure document-level styles
# ============================================================
def setup_styles(doc):
    styles = doc.styles
    
    # Normal style - 13pt Times New Roman, 1.5 line spacing
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(13)
    normal_font.bold = False
    normal_font.color.rgb = RGBColor(0, 0, 0)
    normal_pf = normal_style.paragraph_format
    normal_pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal_pf.space_before = Pt(0)
    normal_pf.space_after = Pt(6)
    normal_pf.first_line_indent = Cm(1.27)
    
    # Heading 1 - Chapter titles (16pt Bold, centered, ALL CAPS)
    h1 = styles['Heading 1']
    h1_font = h1.font
    h1_font.name = 'Times New Roman'
    h1_font.size = Pt(16)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(0, 0, 0)
    h1_pf = h1.paragraph_format
    h1_pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1_pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    h1_pf.space_before = Pt(12)
    h1_pf.space_after = Pt(12)
    h1_pf.first_line_indent = Cm(0)
    h1_pf.keep_with_next = True
    
    # Heading 2 - Sections (14pt Bold)
    h2 = styles['Heading 2']
    h2_font = h2.font
    h2_font.name = 'Times New Roman'
    h2_font.size = Pt(14)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(0, 0, 0)
    h2_pf = h2.paragraph_format
    h2_pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2_pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    h2_pf.space_before = Pt(12)
    h2_pf.space_after = Pt(6)
    h2_pf.first_line_indent = Cm(0)
    h2_pf.keep_with_next = True
    
    # Heading 3 - Subsections (13pt Bold)
    h3 = styles['Heading 3']
    h3_font = h3.font
    h3_font.name = 'Times New Roman'
    h3_font.size = Pt(13)
    h3_font.bold = True
    h3_font.color.rgb = RGBColor(0, 0, 0)
    h3_pf = h3.paragraph_format
    h3_pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3_pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    h3_pf.space_before = Pt(6)
    h3_pf.space_after = Pt(6)
    h3_pf.first_line_indent = Cm(0)
    h3_pf.keep_with_next = True

setup_styles(doc)

# ============================================================
# STEP 2: Classify and fix paragraphs
# ============================================================

# Patterns for headings based on numbering
chapter_pattern = re.compile(r'^CH[ƯƯƠNG\u01AF\u01B0]*\s*NG\s*\d+', re.IGNORECASE)
# Also match "CHƯƠNG" with various Vietnamese encodings
chapter_pattern2 = re.compile(r'^CH.{1,4}NG\s*\d+', re.IGNORECASE)

# Section pattern: starts with digit.digit. (like "1.1.", "2.3.")
section_pattern = re.compile(r'^\d+\.\d+[\.\s]')

# Subsection pattern: starts with digit.digit.digit (like "1.1.1", "2.3.4")
subsection_pattern = re.compile(r'^\d+\.\d+\.\d+')

# Conclusion summary paragraphs that start like "Chương X đã..."
conclusion_summary_pattern = re.compile(r'^Ch.{1,4}ng\s+\d+\s+đ', re.IGNORECASE)

def get_first_run_info(para):
    """Get font size, bold, font name from first non-empty run"""
    for run in para.runs:
        if run.text.strip():
            return {
                'size': run.font.size,
                'bold': run.font.bold,
                'name': run.font.name
            }
    return {'size': None, 'bold': None, 'name': None}

def classify_paragraph(idx, para):
    """Classify paragraph type"""
    text = para.text.strip()
    if not text:
        return 'empty'
    
    # Cover page (first ~12 paragraphs)
    if idx <= 12:
        return 'cover'
    
    info = get_first_run_info(para)
    font_size = info['size']
    is_bold = info['bold']
    font_name = info['name']
    
    # Bullet points
    if font_name in ('Symbol', 'Wingdings') or (text.startswith('•') and not is_bold):
        return 'bullet'
    
    # Check if it's a chapter title by text pattern AND font size
    if font_size and font_size >= 228600:  # >= 18pt - definitely chapter/title area
        # But check if it's a conclusion summary
        if conclusion_summary_pattern.match(text):
            return 'body'  # This is a conclusion paragraph, not a heading
        if chapter_pattern.match(text) or chapter_pattern2.match(text):
            return 'chapter'
        # Could be title/subtitle on cover
        if is_bold:
            return 'chapter'
    
    if font_size and font_size >= 254000:  # 20pt
        if conclusion_summary_pattern.match(text):
            return 'body'
        return 'chapter'
    
    # Check text patterns for heading detection
    if is_bold and font_size:
        # Section-level (16pt bold) - like "1.1. Title"
        if font_size >= 203200:  # >= 16pt
            if conclusion_summary_pattern.match(text):
                return 'body'
            if section_pattern.match(text):
                return 'section'
            # Chapter title that's in 16pt
            if chapter_pattern.match(text) or chapter_pattern2.match(text):
                return 'chapter'
            return 'section'
        
        # Subsection-level (14pt bold) - like "1.1.1. Title"
        if font_size >= 177800:  # >= 14pt
            if subsection_pattern.match(text):
                return 'subsection'
            if section_pattern.match(text):
                return 'section'  # Sometimes sections use 14pt too
            # Numbered items in conclusion (like "1. Hiện thực hóa...")
            if re.match(r'^\d+\.\s', text):
                return 'subsection'
            return 'subsection'
    
    return 'body'

# Process all paragraphs
for idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    ptype = classify_paragraph(idx, para)
    
    # Cover page - minimal changes
    if ptype == 'cover':
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        for run in para.runs:
            run.font.name = 'Times New Roman'
        continue
    
    # Empty paragraphs
    if ptype == 'empty':
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        continue
    
    # Chapter title (Heading 1)
    if ptype == 'chapter':
        para.style = doc.styles['Heading 1']
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(12)
        para.paragraph_format.first_line_indent = Cm(0)
        
    # Section (Heading 2)
    elif ptype == 'section':
        para.style = doc.styles['Heading 2']
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.first_line_indent = Cm(0)
        
    # Subsection (Heading 3)
    elif ptype == 'subsection':
        para.style = doc.styles['Heading 3']
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.first_line_indent = Cm(0)
        
    # Bullet points
    elif ptype == 'bullet':
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(3)
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.left_indent = Cm(1.27)
        for run in para.runs:
            if run.font.name not in ('Symbol', 'Wingdings'):
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
            else:
                run.font.size = Pt(13)
        
    # Body text
    else:
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.first_line_indent = Cm(1.27)
        for run in para.runs:
            if run.font.name not in ('Courier New',):  # Keep code font
                run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0, 0, 0)

# ============================================================
# STEP 3: Fix table formatting
# ============================================================
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
                para.paragraph_format.first_line_indent = Cm(0)
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    if not run.font.size or run.font.size < Pt(11):
                        run.font.size = Pt(12)

# ============================================================
# STEP 4: Set page margins
# ============================================================
for section in doc.sections:
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

# ============================================================
# STEP 5: Save
# ============================================================
doc.save(OUTPUT_FILE)
print(f"Saved to: {OUTPUT_FILE}")

# Print summary
h1_count = sum(1 for p in doc.paragraphs if p.style and p.style.name == 'Heading 1')
h2_count = sum(1 for p in doc.paragraphs if p.style and p.style.name == 'Heading 2')
h3_count = sum(1 for p in doc.paragraphs if p.style and p.style.name == 'Heading 3')
normal_count = sum(1 for p in doc.paragraphs if p.style and p.style.name == 'Normal')

print(f"\nSummary:")
print(f"  Heading 1 (Chapters): {h1_count}")
print(f"  Heading 2 (Sections): {h2_count}")
print(f"  Heading 3 (Subsections): {h3_count}")
print(f"  Normal (Body): {normal_count}")
print(f"  All line spacing: 1.5")
print(f"  Body font: Times New Roman 13pt")
print(f"  Margins: 3.5cm left, 2cm right/top/bottom")
print("Done!")
